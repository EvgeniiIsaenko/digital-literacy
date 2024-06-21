# Лаба №2: Приглашение
from docx import *
from docx.shared import Cm, Inches

document = Document()

document.add_heading("Приглашение на свадьбу", 0)

# TODO: картинка,
p1 = document.add_paragraph(r"Приглашаем Вас, {{ full_name }}, на ")
p1.add_run("свадьбу").bold = True
p1.add_run(". Проводим свадьбу в ")
p1.add_run("чебуречной ").italic = True
p1.add_run(r"по адресу {{ address }}.")

document.add_heading("Что Вам необходимо взять с собой: ")
p2 = document.add_paragraph("алкоголь;", style="List Bullet")
document.add_paragraph("хорошее настроение.", style="List Bullet")

table1_contents = (
    (30, 'Водка'),
    (40, 'Чебуреки'),
    (81, 'Детское шампанское'),
)

table1 = document.add_table(rows=1, cols=2)
cols_table1 = table1.rows[0].cells
cols_table1[0].text = 'Кол-во'
cols_table1[1].text = 'Предмет'
for qty, name in table1_contents:
    rows_table1 = table1.add_row().cells
    rows_table1[0].text = str(qty)
    rows_table1[1].text = name

document.add_paragraph()
document.add_picture("pictures/cheburek.jpg", width=Cm(7.0))

p3 = document.add_paragraph(r"Начало торжества в {{ time }}, {{ date }}")

p4 = document.add_paragraph("Подпись ")
p4.add_run("_________")
p4.add_run().add_picture("pictures/signature.png", width=Cm(2.0))
p4.add_run("_________")

pic2 = document.add_picture("pictures/stamp.jpg", width=Cm(3.0))

document.add_page_break()

document.save("template/invitation_tpl.docx")
